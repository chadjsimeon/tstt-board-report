"""
TSTT Board Report — high-fidelity screenshot Word exporter.

Drives headless Chromium over the live Streamlit dashboard, screenshots every
page (Home + each sidebar page), and assembles the captures into a .docx board
pack — one dashboard page per Word page, under a heading, at full page width.
Pages taller than one printed page are auto-split.

The browser capture, page discovery and slicing are shared with the PowerPoint
exporter (`pptx_screenshot_export`); only the document assembly differs.

Usage (from the project root):
    python tools/docx_screenshot_export.py
    python tools/docx_screenshot_export.py --url http://localhost:8501
    python tools/docx_screenshot_export.py --pages "Consumer Sales,Cash CAPEX"

One-time setup:
    pip install playwright Pillow python-docx
    playwright install chromium
"""

from __future__ import annotations

import argparse
import io
import sys
from datetime import date
from pathlib import Path

try:                                   # imported as `tools.docx_screenshot_export`
    from tools.pptx_screenshot_export import (
        PROJECT_ROOT,
        _launch_streamlit,
        _wait_for_server,
        capture_all,
        discover_pages,
        slice_capture,
    )
except ImportError:                    # run directly: tools/ is on sys.path
    from pptx_screenshot_export import (
        PROJECT_ROOT,
        _launch_streamlit,
        _wait_for_server,
        capture_all,
        discover_pages,
        slice_capture,
    )

# ── Page geometry (US Letter landscape, matching the dashboard's wide layout) ──
PAGE_W_IN   = 11.0
PAGE_H_IN   = 8.5
MARGIN_IN   = 0.45
HDR_FTR_IN  = 0.30        # must stay <= MARGIN_IN, see build_docx
HEADING_IN  = 0.55        # vertical room reserved for the per-page heading

HEADING_PT       = 16     # Heading 1 size
HEADING_AFTER_PT = 6      # space below the heading

PORTRAIT_W_IN = 8.5
PORTRAIT_H_IN = 11.0

BRAND_GREEN = (0x00, 0x87, 0x5A)
INK         = (0x1F, 0x23, 0x28)
MUTED       = (0x5B, 0x66, 0x75)


def _geometry(portrait: bool):
    """Return (page_w, page_h, image_w, image_h, aspect) in inches."""
    pw, ph = (PORTRAIT_W_IN, PORTRAIT_H_IN) if portrait else (PAGE_W_IN, PAGE_H_IN)
    img_w = pw - 2 * MARGIN_IN
    img_h = ph - 2 * MARGIN_IN - HEADING_IN
    return pw, ph, img_w, img_h, img_w / img_h


# ── DOCX assembly ────────────────────────────────────────────────────────────
def require_docx() -> None:
    """Fail fast if python-docx is missing.

    Called before the browser capture as well as at assembly time: capturing every
    page takes 2-3 minutes, and discovering the missing dependency only afterwards
    throws that work away. The interpreter is named explicitly because this project
    is commonly run from a different Python than the one holding `venv/`.
    """
    try:
        import docx  # noqa: F401
    except ImportError:
        raise RuntimeError(
            "python-docx is not installed for this interpreter "
            f"({sys.executable}).\n"
            f'  Install it with:  "{sys.executable}" -m pip install python-docx'
        ) from None


def build_docx(captures, out_path: Path, split: bool, portrait: bool = False,
               period_label: str = "") -> int:
    require_docx()
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    pw, ph, img_w, img_h, aspect = _geometry(portrait)

    doc = Document()

    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT if portrait else WD_ORIENT.LANDSCAPE
    # python-docx does not swap the dimensions for you when changing orientation.
    sec.page_width  = Inches(pw)
    sec.page_height = Inches(ph)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, attr, Inches(MARGIN_IN))
    # Word reserves max(margin, header/footer distance) — the stock 0.5in defaults
    # would silently eat into the text area and push the image onto its own page.
    sec.header_distance = Inches(HDR_FTR_IN)
    sec.footer_distance = Inches(HDR_FTR_IN)

    # Body/Heading styles tuned for a printed board pack.
    #
    # The stock Normal style carries 1.08 line spacing and 8pt space-after, which
    # silently inflates a full-height inline image past the text area and bumps it
    # onto the next page, orphaning its heading. Zero both out.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(*INK)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(HEADING_PT)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(*INK)
    h1.paragraph_format.line_spacing = 1.0
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(HEADING_AFTER_PT)
    h1.paragraph_format.keep_with_next = True

    # ── Cover page ───────────────────────────────────────────────────────────
    def cover_line(text, size, colour, bold=False, after=0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.color.rgb = RGBColor(*colour)
        return p

    # Push the block roughly a third of the way down the page.
    doc.add_paragraph().paragraph_format.space_after = Pt(
        (PORTRAIT_H_IN if portrait else PAGE_H_IN) * 72 * 0.22
    )

    cover_line("TSTT", 40, BRAND_GREEN, bold=True, after=4)
    cover_line("Board Report", 28, INK, bold=True, after=10)
    cover_line("Telecommunications Services of Trinidad and Tobago", 12, MUTED, after=18)
    if period_label:
        cover_line(period_label, 14, BRAND_GREEN, bold=True, after=18)
    cover_line(
        f"Generated {date.today():%d %B %Y}  ·  Confidential — Board Use Only",
        9, MUTED,
    )

    # ── One dashboard page per Word page ─────────────────────────────────────
    n_pages = 0
    for label, png in captures:
        segments = slice_capture(png, split, aspect=aspect)
        for i, seg in enumerate(segments):
            heading = f"{label} (cont. {i + 1})" if i else label
            h = doc.add_heading(heading, level=1)
            # Start each dashboard page on a fresh sheet. Using page_break_before
            # rather than a break paragraph avoids an empty line at the page top.
            h.paragraph_format.page_break_before = True

            bio = io.BytesIO()
            seg.save(bio, format="PNG")
            bio.seek(0)

            sw, sh = seg.size
            if split:
                # Segments already match the page aspect => fill the text width.
                doc.add_picture(bio, width=Inches(img_w))
            else:
                # Fit the whole capture inside the image box, preserving aspect.
                if sw / sh >= aspect:
                    doc.add_picture(bio, width=Inches(img_w))
                else:
                    doc.add_picture(bio, height=Inches(img_h))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            n_pages += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"\n  wrote {n_pages} page(s) -> {out_path}")
    return n_pages


# ── Main ─────────────────────────────────────────────────────────────────────
def main() -> int:
    ap = argparse.ArgumentParser(description="Screenshot-based all-pages Word exporter.")
    ap.add_argument("--url", default=None,
                    help="Reuse a running Streamlit server (e.g. http://localhost:8501). "
                         "If omitted, a dedicated instance is launched.")
    ap.add_argument("--port", type=int, default=8599,
                    help="Port for the auto-launched Streamlit instance (default 8599).")
    ap.add_argument("--page", default=None,
                    help="Optional: capture only this one page slug (default: all pages).")
    ap.add_argument("--pages", default=None,
                    help="Optional: comma-separated page labels/slugs to capture "
                         "(default: all). Takes precedence over --page.")
    ap.add_argument("--focus-month", default=None,
                    help="Focus month to render every page at (e.g. Jun-26). "
                         "Passed to the app as a query param; default: latest month.")
    ap.add_argument("--period-label", default=None,
                    help="Period text for the cover page (e.g. 'June 2026').")
    ap.add_argument("--out", default="exports/TSTT_Board_Report.docx",
                    help="Output .docx path (default: exports/TSTT_Board_Report.docx).")
    ap.add_argument("--viewport-width", type=int, default=1920,
                    help="Browser viewport width in CSS px (default 1920).")
    ap.add_argument("--portrait", action="store_true",
                    help="Portrait pages (default: landscape, which suits the wide layout).")
    ap.add_argument("--no-split", action="store_true",
                    help="One Word page per dashboard page, scaled to fit (no auto-split).")
    ap.add_argument("--headed", action="store_true",
                    help="Show the browser window (debugging).")
    args = ap.parse_args()

    # Check before capturing, not after — see require_docx().
    try:
        require_docx()
    except RuntimeError as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1

    out_path = Path(args.out)
    if not out_path.is_absolute():
        out_path = PROJECT_ROOT / out_path

    pages = discover_pages()
    if args.pages:
        wanted = {p.strip() for p in args.pages.split(",") if p.strip()}
        pages = [(l, s) for l, s in pages if l in wanted or s in wanted]
        if not pages:
            print(f"ERROR: no pages match --pages {args.pages!r}.", file=sys.stderr)
            return 1
    elif args.page:
        pages = [(l, s) for l, s in pages if args.page in (s, l)]
        if not pages:
            print(f"ERROR: no page matches --page {args.page!r}.", file=sys.stderr)
            return 1
    print(f"Pages to capture: {', '.join(l for l, _ in pages)}")

    proc = None
    try:
        if args.url:
            base_url = args.url
            print(f"Using existing server: {base_url}")
            _wait_for_server(base_url)
        else:
            proc = _launch_streamlit(args.port)
            base_url = f"http://localhost:{args.port}"
            _wait_for_server(base_url)
            print("  server ready")

        captures = capture_all(base_url, pages, args.viewport_width, args.headed,
                               args.focus_month)
        if not captures:
            print("ERROR: no pages captured.", file=sys.stderr)
            return 1

        build_docx(captures, out_path, split=not args.no_split,
                   portrait=args.portrait,
                   period_label=args.period_label or args.focus_month or "")
        print("Done.")
        return 0

    finally:
        if proc is not None:
            print("  shutting down Streamlit ...")
            proc.terminate()
            try:
                proc.wait(timeout=10)
            except Exception:                       # noqa: BLE001
                proc.kill()


if __name__ == "__main__":
    raise SystemExit(main())
